"""
Диалог для редактирования сопоставлений полей шаблона с данными из БД
✅ ИСПРАВЛЕНО: Полная синхронизация с schema_only.sql через db_mappings.py
✅ ИСПРАВЛЕНО: Убрано дублирование колонок, используется единый DB_COLUMNS_MAP
✅ ИСПРАВЛЕНО: Корректная обработка "recipients.name" вместо "recipient_name"
✅ ДОБАВЛЕНО: SearchableComboBox для быстрого поиска полей
"""
import os
import tempfile
import json
import re
from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QLabel,
    QTableWidget, QHeaderView, QAbstractItemView, QMessageBox,
    QComboBox, QDialogButtonBox, QGroupBox, QTableWidgetItem, QCompleter
)
from ui_helpers import BaseDialog
from PyQt6.QtCore import Qt, QByteArray
from PyQt6.QtSql import QSqlQuery
from PyQt6.QtGui import QFont
from docx import Document
from composite_field_widget import CompositeFieldWidget
from field_mapping_manager import FieldMappingManager
from searchable_combo import SearchableComboBox

# ✅ ЕДИНЫЙ ИСТОЧНИК ДАННЫХ: Импорт включает get_field_description для корректного поиска
try:
    from db_mappings import TABLE_NAMES_RU, COLUMN_DESCRIPTIONS, DB_COLUMNS_MAP, get_field_description
except ImportError:
    TABLE_NAMES_RU = {}
    COLUMN_DESCRIPTIONS = {}
    DB_COLUMNS_MAP = {}
    def get_field_description(t, c): return c  # Фоллбэк если файл не найден

class MappingEditorDialog(BaseDialog):
    """Отдельное окно для редактирования сопоставлений"""
    def __init__(self, parent=None, krd_id=None, db_connection=None, template_id=None, audit_logger=None):
        super().__init__(parent)
        self.krd_id = krd_id
        self.db = db_connection
        self.template_id = template_id
        self.audit_logger = audit_logger
        self.template_variables = []
        self.db_columns = {}
        self.current_template_id = template_id
        
        # Инициализация модулей
        self.composite_widget = CompositeFieldWidget(self)
        self.mapping_manager = FieldMappingManager(self)
        
        self.setWindowTitle("✏️ Редактирование сопоставлений полей")
        self.setMinimumSize(1100, 750)
        self.setModal(False)
        self.init_ui()
        
        if template_id:
            self.load_field_mappings(template_id)

    def init_ui(self):
        layout = QVBoxLayout()
        layout.setSpacing(10)

        # Заголовок
        title_label = QLabel("🔗 Сопоставление полей шаблона с данными из БД")
        title_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        layout.addWidget(title_label)

        info_label = QLabel("💡 Выберите переменную из шаблона и укажите, какое поле из базы данных должно быть подставлено")
        info_label.setStyleSheet("QLabel { color: #666; padding: 8px; background-color: #f0f0f0; border-radius: 5px; }")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)

        # Таблица сопоставлений
        self.mapping_table = QTableWidget()
        self.mapping_table.setColumnCount(3)
        self.mapping_table.setHorizontalHeaderLabels([
            "Переменная из шаблона",
            "Поле из базы данных",
            "Тип"
        ])
        self.mapping_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.mapping_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        
        header = self.mapping_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Fixed)
        self.mapping_table.setColumnWidth(2, 100)
        layout.addWidget(self.mapping_table)

        # Кнопки управления
        btn_group = QGroupBox("Управление сопоставлениями")
        btn_layout = QHBoxLayout(btn_group)

        add_simple_btn = QPushButton("➕ Добавить простое сопоставление")
        add_simple_btn.setStyleSheet("QPushButton { background-color: #2196F3; color: white; font-weight: bold; padding: 8px 15px; border-radius: 5px; } QPushButton:hover { background-color: #1976D2; }")
        add_simple_btn.clicked.connect(self.add_field_mapping)
        btn_layout.addWidget(add_simple_btn)

        add_composite_btn = QPushButton("🔗 Добавить составное поле")
        add_composite_btn.setStyleSheet("QPushButton { background-color: #9C27B0; color: white; font-weight: bold; padding: 8px 15px; border-radius: 5px; } QPushButton:hover { background-color: #7B1FA2; }")
        add_composite_btn.clicked.connect(self.add_composite_field_mapping)
        btn_layout.addWidget(add_composite_btn)

        delete_btn = QPushButton("🗑️ Удалить выбранное")
        delete_btn.setStyleSheet("QPushButton { background-color: #f44336; color: white; font-weight: bold; padding: 8px 15px; border-radius: 5px; } QPushButton:hover { background-color: #d32f2f; }")
        delete_btn.clicked.connect(self.remove_field_mapping)
        btn_layout.addWidget(delete_btn)
        layout.addWidget(btn_group)

        # Кнопки сохранения и закрытия
        button_box = QDialogButtonBox()
        save_btn = QPushButton("💾 Сохранить сопоставления")
        save_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; padding: 12px 30px; border-radius: 5px; font-size: 13px; } QPushButton:hover { background-color: #45a049; }")
        save_btn.clicked.connect(self.save_and_close)
        cancel_btn = QPushButton("❌ Закрыть без сохранения")
        cancel_btn.setStyleSheet("QPushButton { background-color: #757575; color: white; font-weight: bold; padding: 12px 30px; border-radius: 5px; font-size: 13px; } QPushButton:hover { background-color: #616161; }")
        cancel_btn.clicked.connect(self.reject)
        
        button_box.addButton(save_btn, QDialogButtonBox.ButtonRole.AcceptRole)
        button_box.addButton(cancel_btn, QDialogButtonBox.ButtonRole.RejectRole)
        layout.addWidget(button_box)

        self.setLayout(layout)

        # Загрузка данных
        self.load_template_variables(self.template_id)
        self.load_db_columns()

    def add_field_mapping(self):
        """Добавление простого сопоставления"""
        if not self.current_template_id:
            return QMessageBox.warning(self, "Ошибка", "Шаблон не выбран")
        if not self.template_variables:
            return QMessageBox.warning(self, "Ошибка", "Переменные шаблона не загружены")

        row = self.mapping_table.rowCount()
        self.mapping_table.insertRow(row)
        
        var_combo = QComboBox()
        var_combo.addItems(self.template_variables)
        self.mapping_table.setCellWidget(row, 0, var_combo)
        
        col_combo = self._create_db_column_combo()
        self.mapping_table.setCellWidget(row, 1, col_combo)
        
        type_label = QLabel("Простое")
        type_label.setStyleSheet("color: #666; font-size: 11px; font-weight: bold;")
        self.mapping_table.setCellWidget(row, 2, type_label)
        
        self.mapping_table.resizeRowToContents(row)
        self.mapping_table.selectRow(row)

    def add_composite_field_mapping(self):
        """Добавление составного поля"""
        if not self.current_template_id:
            return QMessageBox.warning(self, "Ошибка", "Выберите шаблон")
        if not self.template_variables:
            return QMessageBox.warning(self, "Ошибка", "Переменные шаблона не загружены")
        
        row = self.mapping_table.rowCount()
        self.composite_widget.add_composite_field_mapping(row)

    def remove_field_mapping(self):
        """Удаление сопоставления"""
        selected_rows = self.mapping_table.selectionModel().selectedRows()
        if not selected_rows:
            return QMessageBox.warning(self, "Внимание", "Выберите сопоставление для удаления")
        row = selected_rows[0].row()
        self.mapping_table.removeRow(row)

    def add_simple_mapping_row(self, row, field_name, db_column, table_name):
        """Добавление простого сопоставления при загрузке из БД"""
        self.mapping_table.insertRow(row)
        
        var_combo = QComboBox()
        var_combo.addItems(self.template_variables)
        var_combo.setCurrentText(field_name)
        self.mapping_table.setCellWidget(row, 0, var_combo)
        
        col_combo = self._create_db_column_combo(db_column)
        self.mapping_table.setCellWidget(row, 1, col_combo)
        
        type_label = QLabel("Простое")
        type_label.setStyleSheet("color: #666; font-size: 11px; font-weight: bold;")
        self.mapping_table.setCellWidget(row, 2, type_label)
        
        self.mapping_table.resizeRowToContents(row)

    def add_composite_mapping_row(self, row, field_name, db_columns_json, table_name):
        """Добавление составного сопоставления при загрузке из БД"""
        self.composite_widget.create_composite_field_row(
            row, field_name, db_columns_json, table_name, self.mapping_table
        )

    def load_field_mappings(self, template_id):
        """Загрузка сопоставлений"""
        self.mapping_table.setRowCount(0)
        self.mapping_manager.load_field_mappings(template_id)

    def _create_db_column_combo(self, selected_column=None):
        """Создание ComboBox с русскими описаниями полей БД и префиксом таблицы"""
        try:
            combo = SearchableComboBox()
            if combo is None: return None
            
            combo.setMaxVisibleItems(50)
            combo.view().setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            combo.view().setMinimumHeight(400)
            combo.setMinimumWidth(350)
            combo.setMaximumWidth(600)
            
            completer = combo.completer()
            if completer:
                completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
                completer.setFilterMode(Qt.MatchFlag.MatchContains)
                completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)

            all_columns = []
            for table_name, columns in self.db_columns.items():
                table_name_ru = TABLE_NAMES_RU.get(table_name, table_name)
                for col in columns:
                    # ✅ ИСПОЛЬЗУЕМ get_field_description для поиска ключа вида "table_column"
                    desc = get_field_description(table_name, col)
                    display_desc = f"[{table_name_ru}] {desc}"
                    all_columns.append((col, display_desc, table_name))

            all_columns.sort(key=lambda x: x[1])

            for col_name, col_description, table_name in all_columns:
                combo.addItem(col_description, f"{table_name}|{col_name}")

            if selected_column:
                # Логика выбора сохраненной колонки
                idx = combo.findData(selected_column)
                if idx < 0 and selected_column:
                    for i in range(combo.count()):
                        item_data = combo.itemData(i)
                        if item_data and item_data.endswith(f"|{selected_column}"):
                            idx = i
                            break
                if idx >= 0: 
                    combo.setCurrentIndex(idx)
            
            return combo
        except Exception as e:
            print(f"❌ Ошибка при создании ComboBox: {e}")
            combo = QComboBox()
            combo.addItem("Ошибка загрузки", None)
            return combo

    def load_template_variables(self, template_id):
        """Загрузка переменных из шаблона DOCX"""
        if not template_id: return
        
        query = QSqlQuery(self.db)
        query.prepare("SELECT template_data FROM krd.document_templates WHERE id = ?")
        query.addBindValue(template_id)
        if not query.exec() or not query.next():
            self.template_variables = []
            return

        data = query.value(0)
        template_bytes = bytes(data) if isinstance(data, QByteArray) else (bytes(data) if data else b'')
        if not template_bytes:
            self.template_variables = []
            return

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(template_bytes)
            tmp_path = tmp.name
            try:
                doc = Document(tmp_path)
                vars_set = set()
                for para in doc.paragraphs:
                    vars_set.update(re.findall(r'\{\{([^{}]+)\}\}', para.text))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                vars_set.update(re.findall(r'\{\{([^{}]+)\}\}', para.text))
                self.template_variables = sorted(vars_set)
            except Exception as e:
                print(f"❌ Ошибка загрузки переменных: {e}")
                self.template_variables = ["surname", "name", "patronymic"]
            finally:
                try: os.unlink(tmp_path)
                except: pass

    def load_db_columns(self):
        """✅ ЗАМЕНА: Используем единый словарь из db_mappings.py"""
        self.db_columns = DB_COLUMNS_MAP

    def save_and_close(self):
        """Сохранение и закрытие"""
        if not self.current_template_id:
            return QMessageBox.critical(self, "Ошибка", "Шаблон не выбран")
        
        if self.mapping_manager.save_field_mappings(self.current_template_id):
            QMessageBox.information(self, "Успех", "✅ Сопоставления успешно сохранены!")
            self.accept()
        else:
            QMessageBox.critical(self, "Ошибка", "❌ Не удалось сохранить сопоставления")