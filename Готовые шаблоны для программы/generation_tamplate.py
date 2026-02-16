
"""
Скрипт для генерации шаблона запроса в миграционную службу
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os


def create_migration_request_template():
    """
    Создание шаблона запроса в миграционную службу с переменными
    """
    # Создаем новый документ
    doc = Document()
    
    # Устанавливаем стиль по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = 12
    
    # Адресат (выравнивание по правому краю)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run('Начальнику отдела по вопросам миграции\n')
    run.bold = True
    
    run = paragraph.add_run('УМВД России г. Абакан\n')
    run.bold = True
    
    run = paragraph.add_run('{{recipient_fio}}\n')
    
    run = paragraph.add_run('{{recipient_address}}\n')
    
    run = paragraph.add_run('{{recipient_phone}}\n')
    
    # Пустая строка
    doc.add_paragraph()
    
    # Основной текст запроса
    paragraph = doc.add_paragraph()
    
    run = paragraph.add_run('Военной комендатурой (гарнизона, 3 разряда) (г. Абакан) проводятся разыскные мероприятия в отношении военнослужащего по контракту войсковой части 42038 рядового ')
    run.bold = False
    
    run = paragraph.add_run('{{surname}} {{name}} {{patronymic}}')
    run.bold = True
    
    run = paragraph.add_run(', {{birth_date}} года рождения, уроженца {{birth_place_town}}, зарегистрированного (проживающего) по адресу: {{registration_address}}, паспорт {{passport_series}} {{passport_number}}, выданный {{passport_issue_date}} {{passport_issued_by}}.')
    
    # Пустая строка
    doc.add_paragraph()
    
    # Продолжение основного текста
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Руководствуясь, положениями межведомственного приказа от 24 октября 2023 года № 92/1/12533/113/22505/1/15-МВС «О мерах по активизации розыска лиц, совершивших преступления против порядка пребывания на военной службе» прошу Вас оказать содействие в розыске указанного военнослужащего, для чего поставить последнего на сторожевой контроль, а также сообщить в адрес военной комендатуры актуальные сведения о его регистрации по месту жительства, месте пребывания, паспортные данные.')
    
    # Пустая строка
    doc.add_paragraph()
    
    # Продолжение
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('В случае установления местонахождения военнослужащего поршу сообщить в военную комендатуру (гарнизона, 3 разряда) (г. Абакан).')
    
    # Пустая строка
    doc.add_paragraph()
    
    # Адрес для ответа
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Ответ прошу направить по адресу: {{response_address}}. ')
    run.bold = True
    
    run = paragraph.add_run('Телефон для связи (факс) {{contact_phone}}.')
    
    # Пустые строки перед подписью
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run('Временно исполняющий обязанности\n')
    run.bold = True
    
    run = paragraph.add_run('военного коменданта военной комендатуры\n')
    run.bold = True
    
    run = paragraph.add_run('(гарнизона, 3 разряда) (г. Абакан)\n')
    run.bold = True
    
    run = paragraph.add_run('лейтенант юстиции\n')
    run.bold = True
    
    run = paragraph.add_run('{{signatory_name}}')
    run.bold = True
    
    # Сохраняем документ
    filename = 'Шаблон_запроса_в_миграцию.docx'
    doc.save(filename)
    print(f"Шаблон сохранен как {filename}")
    
    return filename


def create_field_mappings_for_template(template_id):
    """
    Создание сопоставления полей для шаблона запроса в миграцию
    """
    mappings = [
        {
            'field_name': '{{recipient_fio}}',
            'db_column': 'recipient_fio',
            'table_name': 'migration_request_recipients',
            'description': 'ФИО получателя запроса'
        },
        {
            'field_name': '{{recipient_address}}',
            'db_column': 'address',
            'table_name': 'migration_request_recipients',
            'description': 'Адрес получателя запроса'
        },
        {
            'field_name': '{{recipient_phone}}',
            'db_column': 'phone',
            'table_name': 'migration_request_recipients',
            'description': 'Телефон получателя запроса'
        },
        {
            'field_name': '{{surname}}',
            'db_column': 'surname',
            'table_name': 'social_data',
            'description': 'Фамилия военнослужащего'
        },
        {
            'field_name': '{{name}}',
            'db_column': 'name',
            'table_name': 'social_data',
            'description': 'Имя военнослужащего'
        },
        {
            'field_name': '{{patronymic}}',
            'db_column': 'patronymic',
            'table_name': 'social_data',
            'description': 'Отчество военнослужащего'
        },
        {
            'field_name': '{{birth_date}}',
            'db_column': 'birth_date',
            'table_name': 'social_data',
            'description': 'Дата рождения военнослужащего'
        },
        {
            'field_name': '{{birth_place_town}}',
            'db_column': 'birth_place_town',
            'table_name': 'social_data',
            'description': 'Место рождения (город)'
        },
        {
            'field_name': '{{registration_address}}',
            'db_column': 'registration_address',
            'table_name': 'addresses',
            'description': 'Адрес регистрации'
        },
        {
            'field_name': '{{passport_series}}',
            'db_column': 'passport_series',
            'table_name': 'social_data',
            'description': 'Серия паспорта'
        },
        {
            'field_name': '{{passport_number}}',
            'db_column': 'passport_number',
            'table_name': 'social_data',
            'description': 'Номер паспорта'
        },
        {
            'field_name': '{{passport_issue_date}}',
            'db_column': 'passport_issue_date',
            'table_name': 'social_data',
            'description': 'Дата выдачи паспорта'
        },
        {
            'field_name': '{{passport_issued_by}}',
            'db_column': 'passport_issued_by',
            'table_name': 'social_data',
            'description': 'Кем выдан паспорт'
        },
        {
            'field_name': '{{response_address}}',
            'db_column': 'response_address',
            'table_name': 'organizations',
            'description': 'Адрес для ответа'
        },
        {
            'field_name': '{{contact_phone}}',
            'db_column': 'contact_phone',
            'table_name': 'organizations',
            'description': 'Контактный телефон'
        },
        {
            'field_name': '{{signatory_name}}',
            'db_column': 'signatory_name',
            'table_name': 'organizations',
            'description': 'ФИО подписывающего лица'
        }
    ]
    
    print("SQL-запросы для добавления сопоставления полей:")
    for mapping in mappings:
        sql = f"""
INSERT INTO krd.field_mappings (template_id, field_name, db_column, table_name, description)
VALUES ({template_id}, '{mapping['field_name']}', '{mapping['db_column']}', '{mapping['table_name']}', '{mapping['description']}');
"""
        print(sql)
    
    return mappings


if __name__ == "__main__":
    # Создаем шаблон
    template_filename = create_migration_request_template()
    
    # Пример использования сопоставления полей
    # (это нужно будет выполнить после добавления шаблона в базу данных)
    print("\nДля добавления сопоставления полей в базу данных используйте:")
    print("# template_id = ID шаблона в базе данных")
    print("# create_field_mappings_for_template(template_id)")