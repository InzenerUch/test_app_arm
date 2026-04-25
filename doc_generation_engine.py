"""
Движок генерации документов из Word-шаблонов
✅ С ПОЛНОЙ ДИАГНОСТИКОЙ ТЕКСТА (ДО/ПОСЛЕ)
✅ ВЫВОДИТ ВЕСЬ ТЕКСТ ШАБЛОНА И ГОТОВОГО ДОКУМЕНТА
"""
import os
import tempfile
import json
import re
from docx.shared import Pt, Cm
from docx import Document
from PyQt6.QtSql import QSqlQuery
from PyQt6.QtCore import QByteArray, QDate
import traceback

class DocGenerationEngine:
    def __init__(self, db, krd_id, audit_logger=None):
        self.db = db
        self.krd_id = krd_id
        self.audit_logger = audit_logger
        self.tables_with_selection = {"addresses", "service_places", "soch_episodes"}
        self.db_columns_map = {}
        self.placeholder_pattern = re.compile(r'\{\{([^{}]+)\}\}')
        self.debug_mode = True  # ✅ ВКЛЮЧИТЬ ДИАГНОСТИКУ

    def set_columns_map(self, cols_map):
        self.db_columns_map = cols_map

    def _log(self, message, level="INFO"):
        """Вспомогательный метод для логирования"""
        if self.debug_mode:
            prefix = {
                "INFO": "📝",
                "WARN": "⚠️",
                "ERROR": "❌",
                "SUCCESS": "✅",
                "DATA": "📊",
                "TEXT": "📄"
            }.get(level, "•")
            print(f"{prefix} [{level}] {message}")

    def _extract_all_text(self, doc):
        """Извлекает ВЕСЬ текст из документа (абзацы + таблицы + колонтитулы)"""
        all_text = []
        
        # Основной текст
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                all_text.append(f"[Абзац {i}] {para.text}")
        
        # Таблицы
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.text.strip():
                            all_text.append(f"[Таблица {i}, Ячейка {j},{k}] {para.text}")
        
        # Колонтитулы
        for i, section in enumerate(doc.sections):
            for para in section.header.paragraphs:
                if para.text.strip():
                    all_text.append(f"[Шапка {i}] {para.text}")
            for para in section.footer.paragraphs:
                if para.text.strip():
                    all_text.append(f"[Подвал {i}] {para.text}")
        
        return all_text

    def build_context(self, template_id, selections):
        """Сбор контекста с диагностикой"""
        self._log("=" * 80, "INFO")
        self._log(f"НАЧАЛО СБОРКИ КОНТЕКСТА (template_id={template_id}, krd_id={self.krd_id})", "INFO")
        self._log("=" * 80, "INFO")
        
        context = {}
        query = QSqlQuery(self.db)
        query.prepare("""
            SELECT field_name, db_column, table_name, db_columns, is_composite
            FROM krd.field_mappings
            WHERE template_id = ?
        """)
        query.addBindValue(template_id)
        
        if not query.exec():
            self._log(f"Ошибка загрузки маппингов: {query.lastError().text()}", "ERROR")
            return context

        mappings_count = 0
        while query.next():
            field_name = query.value(0).strip('{} ')
            db_column = query.value(1)
            table_name = query.value(2)
            db_columns_json = query.value(3)
            is_composite = query.value(4) or False

            try:
                if is_composite and db_columns_json:
                    value = self._get_composite_value(table_name, db_columns_json, selections)
                    source = f"COMPOSITE({table_name})"
                else:
                    selected_id = selections.get(table_name)
                    if selected_id and table_name in self.tables_with_selection:
                        value = self._get_value_from_record(table_name, db_column, selected_id)
                        source = f"{table_name}.id={selected_id}"
                    else:
                        value = self._get_value_from_social_data(db_column)
                        source = "social_data"
                
                if value is not None and value != "":
                    context[field_name] = value
                    self._log(f"{{{field_name}}} = '{value}' (источник: {source})", "DATA")
                    mappings_count += 1
                else:
                    self._log(f"{{{field_name}}} = ПУСТО (источник: {source})", "WARN")
            except Exception as e:
                self._log(f"Ошибка получения {{{field_name}}}: {e}", "ERROR")
        
        self._log("=" * 80, "INFO")
        self._log(f"КОНТЕКСТ СОБРАН: {mappings_count} переменных из {query.size()}", "SUCCESS")
        self._log(f"ВСЕГО переменных в контексте: {len(context)}", "DATA")
        self._log("=" * 80, "INFO")
        return context

    def _get_composite_value(self, table_hint, db_columns_json, selections):
        try:
            db_columns = json.loads(db_columns_json) if isinstance(db_columns_json, str) else db_columns_json
            if not db_columns: return None
            parts = []
            for col_info in db_columns:
                col = col_info.get('column')
                sep = col_info.get('separator', '')
                if not col: continue
                t = self._get_table_by_column(col) or table_hint
                sid = selections.get(t)
                val = self._get_value_from_record(t, col, sid) if (sid and t in self.tables_with_selection) else self._get_value_from_social_data(col)
                if val:
                    parts.append(str(val))
                    if sep: parts.append(sep)
            if parts and parts[-1] in [', ', ' ', '; ', ': ', ' - ']: parts.pop()
            return ''.join(parts) if parts else None
        except Exception as e:
            self._log(f"Ошибка составного поля: {e}", "ERROR")
            return None

    def _get_table_by_column(self, col):
        for t, cols in self.db_columns_map.items():
            if col in cols: return t
        return None

    def _get_value_from_record(self, table, col, rid):
        if not re.match(r'^\w+$', table) or not re.match(r'^\w+$', col): return ""
        q = QSqlQuery(self.db)
        q.prepare(f"SELECT {col} FROM krd.{table} WHERE id = ?")
        q.addBindValue(rid)
        if q.exec() and q.next(): return self._format_value(q.value(0))
        return ""

    def _get_value_from_social_data(self, col):
        if not re.match(r'^\w+$', col): return ""
        q = QSqlQuery(self.db)
        q.prepare(f"SELECT {col} FROM krd.social_data WHERE krd_id = ? ORDER BY id DESC LIMIT 1")
        q.addBindValue(self.krd_id)
        if q.exec() and q.next(): return self._format_value(q.value(0))
        return ""

    def _format_value(self, val):
        if val is None: return ""
        if hasattr(val, 'getDate'):
            y, m, d = val.getDate()
            return f"{d:02d}.{m:02d}.{y}"
        return str(val)

    def apply_to_docx(self, template_bytes, context):
        """Применение контекста к шаблону DOCX с ПОЛНОЙ диагностикой текста"""
        self._log("=" * 80, "INFO")
        self._log("НАЧАЛО ГЕНЕРАЦИИ ДОКУМЕНТА", "INFO")
        self._log("=" * 80, "INFO")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(template_bytes)
            template_path = tmp.name
        
        try:
            self._log(f"Загрузка шаблона: {len(template_bytes)} байт", "DATA")
            doc = Document(template_path)
            
            # ═══════════════════════════════════════════════════════════════
            # ✅ 1. ВЫВОД ПОЛНОГО ТЕКСТА ШАБЛОНА (ДО ГЕНЕРАЦИИ)
            # ═══════════════════════════════════════════════════════════════
            self._log("\n" + "=" * 80, "TEXT")
            self._log("📄 ТЕКСТ ШАБЛОНА (ДО ГЕНЕРАЦИИ)", "TEXT")
            self._log("=" * 80, "TEXT")
            
            template_text_lines = self._extract_all_text(doc)
            for line in template_text_lines:
                self._log(line, "TEXT")
            
            if not template_text_lines:
                self._log("⚠️ ШАБЛОН ПУСТОЙ ИЛИ НЕ СОДЕРЖИТ ТЕКСТА!", "ERROR")
            
            self._log(f"\nВСЕГО строк текста в шаблоне: {len(template_text_lines)}", "DATA")
            self._log("=" * 80, "TEXT")
            
            # 🔍 Поиск переменных в шаблоне
            template_vars = set()
            self._log("\n🔍 ПОИСК ПЕРЕМЕННЫХ В ШАБЛОНЕ:", "INFO")
            
            for paragraph in doc.paragraphs:
                found = self.placeholder_pattern.findall(paragraph.text)
                if found:
                    template_vars.update(found)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            found = self.placeholder_pattern.findall(paragraph.text)
                            if found:
                                template_vars.update(found)
            
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    found = self.placeholder_pattern.findall(paragraph.text)
                    if found:
                        template_vars.update(found)
                for paragraph in section.footer.paragraphs:
                    found = self.placeholder_pattern.findall(paragraph.text)
                    if found:
                        template_vars.update(found)
            
            self._log(f"ВСЕГО найдено переменных в шаблоне: {len(template_vars)}", "SUCCESS")
            self._log(f"Список: {', '.join(sorted(template_vars))}", "DATA")
            
            # 🔍 Проверка соответствия
            context_vars = set(context.keys())
            missing_in_context = template_vars - context_vars
            unused_in_template = context_vars - template_vars
            
            if missing_in_context:
                self._log(f"\n⚠️ ПЕРЕМЕННЫЕ В ШАБЛОНЕ, НО НЕТ В КОНТЕКСТЕ ({len(missing_in_context)}):", "WARN")
                for var in sorted(missing_in_context):
                    self._log(f"   {{{var}}}", "WARN")
            
            if unused_in_template:
                self._log(f"\n⚠️ ПЕРЕМЕННЫЕ В КОНТЕКСТЕ, НО НЕТ В ШАБЛОНЕ ({len(unused_in_template)}):", "WARN")
                for var in sorted(unused_in_template):
                    self._log(f"   {{{var}}}", "WARN")
            
            # 🔄 Замена переменных
            self._log("\n🔄 ЗАМЕНА ПЕРЕМЕННЫХ:", "INFO")
            replacements = 0
            replacement_stats = {}
            
            for paragraph in doc.paragraphs:
                count = self._replace_in_paragraph(paragraph, context, replacement_stats)
                replacements += count
                
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            count = self._replace_in_paragraph(paragraph, context, replacement_stats)
                            replacements += count
                            
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    count = self._replace_in_paragraph(paragraph, context, replacement_stats)
                    replacements += count
                for paragraph in section.footer.paragraphs:
                    count = self._replace_in_paragraph(paragraph, context, replacement_stats)
                    replacements += count
            
            self._log(f"\nВСЕГО заменено: {replacements}", "SUCCESS")
            
            if replacement_stats:
                self._log("\n📊 СТАТИСТИКА ЗАМЕН ПО ПЕРЕМЕННЫМ:", "DATA")
                for var, count in sorted(replacement_stats.items(), key=lambda x: x[1], reverse=True):
                    self._log(f"   {{{var}}}: {count} раз(а)", "DATA")
            
            # 💾 Сохранение
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as out_file:
                doc.save(out_file.name)
                output_path = out_file.name
            
            output_size = os.path.getsize(output_path)
            self._log(f"\n💾 СОХРАНЕНИЕ РЕЗУЛЬТАТА:", "SUCCESS")
            self._log(f"  Путь: {output_path}", "DATA")
            self._log(f"  Размер: {output_size} байт", "DATA")
            
            # ═══════════════════════════════════════════════════════════════
            # ✅ 2. ВЫВОД ПОЛНОГО ТЕКСТА ГОТОВОГО ДОКУМЕНТА (ПОСЛЕ ГЕНЕРАЦИИ)
            # ═══════════════════════════════════════════════════════════════
            self._log("\n" + "=" * 80, "TEXT")
            self._log("📄 ТЕКСТ ГОТОВОГО ДОКУМЕНТА (ПОСЛЕ ГЕНЕРАЦИИ)", "TEXT")
            self._log("=" * 80, "TEXT")
            
            result_doc = Document(output_path)
            result_text_lines = self._extract_all_text(result_doc)
            for line in result_text_lines:
                self._log(line, "TEXT")
            
            if not result_text_lines:
                self._log("⚠️ ГОТОВЫЙ ДОКУМЕНТ ПУСТОЙ!", "ERROR")
            
            self._log(f"\nВСЕГО строк текста в готовом документе: {len(result_text_lines)}", "DATA")
            self._log("=" * 80, "TEXT")
            
            # 🔍 Проверка оставшихся переменных
            remaining_vars = set()
            for paragraph in result_doc.paragraphs:
                found = self.placeholder_pattern.findall(paragraph.text)
                if found:
                    remaining_vars.update(found)
            
            for table in result_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            found = self.placeholder_pattern.findall(paragraph.text)
                            if found:
                                remaining_vars.update(found)
            
            if remaining_vars:
                self._log(f"\n⚠️ В ИТОГОВОМ ДОКУМЕНТЕ ОСТАЛИСЬ ПЕРЕМЕННЫЕ ({len(remaining_vars)}):", "WARN")
                for var in sorted(remaining_vars):
                    self._log(f"   {{{var}}}", "WARN")
            else:
                self._log("\n✅ Все переменные заменены успешно", "SUCCESS")
            
            # ═══════════════════════════════════════════════════════════════
            # ✅ 3. СРАВНЕНИЕ ДО/ПОСЛЕ
            # ═══════════════════════════════════════════════════════════════
            self._log("\n" + "=" * 80, "INFO")
            self._log("📊 СРАВНЕНИЕ ДО/ПОСЛЕ", "INFO")
            self._log("=" * 80, "INFO")
            self._log(f"Строк в шаблоне: {len(template_text_lines)}", "DATA")
            self._log(f"Строк в готовом документе: {len(result_text_lines)}", "DATA")
            self._log(f"Заменено переменных: {replacements}", "DATA")
            self._log(f"Осталось переменных: {len(remaining_vars)}", "DATA")
            self._log("=" * 80, "INFO")
            self._log("ГЕНЕРАЦИЯ ЗАВЕРШЕНА", "SUCCESS")
            self._log("=" * 80, "INFO")
            
            return output_path, replacements
            
        finally:
            if os.path.exists(template_path):
                os.unlink(template_path)

    def _replace_in_paragraph(self, paragraph, context, stats_dict=None):
        """Замена с диагностикой"""
        original_text = paragraph.text
        if not original_text:
            return 0
        
        new_text = original_text
        replacements = 0
        
        for var_name, value in context.items():
            placeholder = f"{{{{{var_name}}}}}"
            if placeholder in new_text:
                count = new_text.count(placeholder)
                replacements += count
                new_text = new_text.replace(placeholder, str(value))
                
                if stats_dict is not None:
                    stats_dict[var_name] = stats_dict.get(var_name, 0) + count
        
        if replacements > 0:
            self._log(f"  Замена в абзаце: {replacements} переменных", "DATA")
            self._log(f"    ДО:  '{original_text[:200]}{'...' if len(original_text) > 200 else ''}'", "DATA")
            self._log(f"    ПОСЛЕ: '{new_text[:200]}{'...' if len(new_text) > 200 else ''}'", "DATA")
            
            if new_text == original_text:
                self._log(f"    ⚠️ ВНИМАНИЕ: Текст НЕ изменился!", "ERROR")
            
            if paragraph.runs:
                first_run = paragraph.runs[0]
                saved_style = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_name': first_run.font.name if first_run.font else None,
                    'font_size': first_run.font.size.pt if first_run.font and first_run.font.size else None,
                }
                
                first_run.text = new_text
                first_run.bold = saved_style['bold']
                first_run.italic = saved_style['italic']
                first_run.underline = saved_style['underline']
                if first_run.font:
                    if saved_style['font_name']: first_run.font.name = saved_style['font_name']
                    if saved_style['font_size']: first_run.font.size = Pt(saved_style['font_size'])
                
                for i in range(len(paragraph.runs) - 1, 0, -1):
                    run = paragraph.runs[i]
                    r_elem = run._element
                    if r_elem in paragraph._element:
                        try:
                            paragraph._element.remove(r_elem)
                        except Exception:
                            pass
            else:
                paragraph.clear()
                run = paragraph.add_run(new_text)
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
        
        return replacements

    def save_to_database(self, request_type_id, recipient_id, issue_number, document_bytes):
        q = QSqlQuery(self.db)
        q.prepare("""
            INSERT INTO krd.outgoing_requests 
            (krd_id, request_type_id, recipient_id, issue_date, issue_number, document_data)
            VALUES (?, ?, ?, CURRENT_DATE, ?, ?)
        """)
        q.addBindValue(self.krd_id)
        q.addBindValue(request_type_id)
        q.addBindValue(recipient_id)
        q.addBindValue(issue_number)
        q.addBindValue(QByteArray(document_bytes))
        
        if not q.exec():
            raise Exception(f"Ошибка БД: {q.lastError().text()}")
            
        req_id = q.lastInsertId()
        if self.audit_logger:
            self.audit_logger.log_action('REQUEST_CREATE', 'outgoing_requests', req_id, self.krd_id, f'Создан запрос №{issue_number}')
        return req_id

    def generate_issue_number(self):
        q = QSqlQuery(self.db)
        q.prepare("SELECT COUNT(*) FROM krd.outgoing_requests WHERE krd_id = ? AND issue_date = CURRENT_DATE")
        q.addBindValue(self.krd_id)
        cnt = 1
        if q.exec() and q.next():
            cnt = (q.value(0) or 0) + 1
        return f"КРД-{self.krd_id}/З-{cnt}"