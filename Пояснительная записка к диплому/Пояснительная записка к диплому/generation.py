import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

def fix_citations_in_report(input_path: str, output_path: str):
    doc = Document(input_path)

    # 🔍 Карта внедрения: (уникальная фраза из текста, правильная ссылка)
    # Фразы подобраны так, чтобы совпадать только с нужными абзацами
    CITATION_MAP = [
        ("Федерального закона № 152-ФЗ «О персональных данных»", "[2]"),
        ("Федеральным законом «О военной полиции и органах военной юстиции в ВС РФ»", "[1]"),
        ("приказам ФСТЭК", "[4], [5]"),
        ("требования информационной безопасности устанавливают", "[5], [25]"),
        ("ГОСТ 34.602-89", "[9]"),
        ("гибкого управления проектами", "[10], [11]"),
        ("обработки конфиденциальных данных", "[12], [13]"),
        ("криптографическое хеширование паролей", "[24]"),
        ("надёжность алгоритмов", "[15], [31]"),
        ("аномалий обновления", "[17], [18]"),
        ("сквозного аудита", "[19]"),
        ("производительность запросов", "[20], [31]"),
        ("парадигмы и MVC", "[21]"),
        ("прерывания экспорта", "[32]"),
        ("соблюдением паттерна MVC", "[23]"),
        ("импортозамещения", "[28], [29]"),
        ("3 нормальной форме", "[17]"),
        ("устойчивость к инъекционным атакам", "[26]"),
        ("Процесс верификации реализован", "[27]")
    ]

    # Регулярка для удаления старых черновых ссылок вида [1], [4, 5] и т.д.
    OLD_REF_RE = re.compile(r'\s*\[\d+(?:,\s*\d+)*\]\s*')

    processed_count = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        text = para.text

        # 1️⃣ Удаляем старые ссылки, чтобы избежать дублирования
        text = OLD_REF_RE.sub(' ', text)
        text = re.sub(r'\s+', ' ', text)  # нормализуем пробелы

        # 2️⃣ Вставляем маркеры новых ссылок после целевых фраз
        for phrase, ref in CITATION_MAP:
            if phrase in text:
                # Заменяем только первое вхождение фразы в абзаце
                text = text.replace(phrase, phrase + f" __CIT_{ref}__", 1)
                processed_count += 1

        # 3️⃣ Пересобираем параграф, сохраняя стиль и выравнивание
        style_name = para.style.name
        alignment = para.alignment
        para.clear()
        para.style = style_name
        if alignment:
            para.alignment = alignment

        # 4️⃣ Разбиваем текст по маркерам и создаём Run-объекты
        parts = re.split(r'(__CIT_.*?__)', text)
        for part in parts:
            if part.startswith("__CIT_") and part.endswith("__"):
                # Извлекаем чистую ссылку [X]
                clean_ref = f"[{part[6:-2]}]"
                run = para.add_run(clean_ref)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.bold = True
                run.font.size = Pt(10)  # немного уменьшаем для аккуратности
            else:
                para.add_run(part)

    doc.save(output_path)
    print(f"✅ Обработано параграфов: {len(doc.paragraphs)}")
    print(f"📌 Вставлено ссылок: {processed_count}")
    print(f"💾 Результат сохранён: {output_path}")

if __name__ == "__main__":
    INPUT_FILE = "Отчет о преддипломной практике.docx"
    OUTPUT_FILE = "Отчет_практика_исправленный.docx"
    fix_citations_in_report(INPUT_FILE, OUTPUT_FILE)